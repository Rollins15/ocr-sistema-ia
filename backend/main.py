# =============================================================================
# BACKEND: OCR API com FastAPI
# Ficheiro: backend/main.py
# Executar: uvicorn main:app --reload --host 0.0.0.0 --port 8000
# =============================================================================
#
# ENDPOINTS:
#   POST /ocr/image   → recebe imagem, extrai texto geral, retorna JSON
#   POST /ocr/quiz    → recebe imagem, extrai perguntas A/B/C/D, retorna JSON
#   POST /ocr/folha   → folha azul UCM only (fase=codigo|disciplinas|opcoes)
#   GET  /health      → verificar se o servidor está a funcionar
#
# INSTALAR:
#   pip install fastapi uvicorn python-multipart pytesseract
#               opencv-python pillow numpy python-docx
# =============================================================================

from fastapi import FastAPI, File, Form, UploadFile, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse
import cv2
import numpy as np
import pytesseract
from PIL import Image
import io
import re
import os
import json
import datetime
import logging
import hashlib
import uuid
from typing import Any, Dict, List, Optional, Tuple
from pathlib import Path

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

logger = logging.getLogger("uvicorn.error")

# Configurar Tesseract (ajustar caminho se necessário)
caminhos_tesseract = [
    r'C:\Program Files\Tesseract-OCR\tesseract.exe',
    r'C:\Program Files (x86)\Tesseract-OCR\tesseract.exe',
    r'C:\Users\hp\AppData\Local\Programs\Tesseract-OCR\tesseract.exe',
]
for c in caminhos_tesseract:
    if os.path.exists(c):
        pytesseract.pytesseract.tesseract_cmd = c
        break

# =============================================================================
# INICIALIZAR APP
# =============================================================================
app = FastAPI(
    title="OCR API",
    description="API para extração de texto em imagens com OCR",
    version="1.0.0"
)

# Permitir CORS para React (web) e React Native (mobile)
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],       # em produção: especificar origens
    allow_methods=["*"],
    allow_headers=["*"],
)

# Pasta para guardar resultados
os.makedirs("resultados", exist_ok=True)


# =============================================================================
# FUNÇÕES DE PRÉ-PROCESSAMENTO
# =============================================================================

def bytes_to_cv2(image_bytes: bytes):
    """Converte bytes recebidos via HTTP para imagem OpenCV."""
    nparr = np.frombuffer(image_bytes, np.uint8)
    return cv2.imdecode(nparr, cv2.IMREAD_COLOR)


def pre_processar(imagem, metodo="automatico"):
    """Aplica pré-processamento para melhorar o OCR."""
    cinza = cv2.cvtColor(imagem, cv2.COLOR_BGR2GRAY)
    h, w = cinza.shape
    # Ampliar 1.5x para melhorar leitura
    cinza = cv2.resize(cinza, (int(w*1.5), int(h*1.5)), interpolation=cv2.INTER_CUBIC)

    if metodo == "simples":
        return cinza
    elif metodo == "contraste":
        clahe = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(8,8))
        return clahe.apply(cinza)
    elif metodo == "adaptativo":
        suavizada = cv2.GaussianBlur(cinza, (3,3), 0)
        return cv2.adaptiveThreshold(suavizada, 255,
                   cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY, 21, 10)
    elif metodo == "otsu":
        denoised = cv2.fastNlMeansDenoising(cinza, h=10)
        _, binarizada = cv2.threshold(denoised, 0, 255,
                           cv2.THRESH_BINARY + cv2.THRESH_OTSU)
        return binarizada
    else:  # automatico
        clahe = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(8,8))
        return clahe.apply(cinza)


def extrair_melhor_texto(imagem):
    """
    Testa vários métodos de pré-processamento e retorna
    o que extrai mais texto.
    """
    metodos = ["simples", "contraste", "adaptativo", "otsu"]
    melhor_texto = ""
    max_chars = 0

    for metodo in metodos:
        try:
            processada = pre_processar(imagem, metodo)
            img_pil = Image.fromarray(processada)
            texto = pytesseract.image_to_string(
                img_pil,
                lang="por+eng",
                config="--psm 3 --oem 3"
            ).strip()
            if len(texto) > max_chars:
                max_chars = len(texto)
                melhor_texto = texto
        except Exception:
            try:
                processada = pre_processar(imagem, metodo)
                img_pil = Image.fromarray(processada)
                texto = pytesseract.image_to_string(
                    img_pil, lang="eng", config="--psm 3 --oem 3"
                ).strip()
                if len(texto) > max_chars:
                    max_chars = len(texto)
                    melhor_texto = texto
            except Exception:
                pass

    return melhor_texto


# =============================================================================
# PARSER DE QUIZ
# =============================================================================

def parsear_quiz(texto: str) -> list:
    """
    Analisa o texto extraído e organiza em estrutura de quiz.

    Formato esperado na imagem:
        1. Qual é a capital de Moçambique?
        A) Maputo
        B) Beira
        C) Nampula
        D) Tete

    Retorna lista de dicionários:
        [{ "id": 1, "texto": "...", "opcoes": {"A":"...","B":"...","C":"...","D":"..."} }]
    """
    perguntas = []
    linhas = [l.strip() for l in texto.split('\n') if l.strip()]

    pergunta_atual = None
    opcoes_atuais = {}
    num_atual = None

    # Padrões para detectar perguntas e opções
    padrao_pergunta = re.compile(r'^(\d+)[.\)]\s+(.+)')
    padrao_opcao    = re.compile(r'^([A-Da-d])[.\)]\s+(.+)')

    for linha in linhas:
        match_p = padrao_pergunta.match(linha)
        match_o = padrao_opcao.match(linha)

        if match_p:
            # Guardar pergunta anterior se existir
            if pergunta_atual and opcoes_atuais:
                perguntas.append({
                    "id": num_atual,
                    "texto": pergunta_atual,
                    "opcoes": opcoes_atuais
                })
            # Nova pergunta
            num_atual = int(match_p.group(1))
            pergunta_atual = match_p.group(2).strip()
            opcoes_atuais = {}

        elif match_o and pergunta_atual is not None:
            letra = match_o.group(1).upper()
            opcoes_atuais[letra] = match_o.group(2).strip()

        elif pergunta_atual and not match_o:
            # Continuação da pergunta (texto em múltiplas linhas)
            pergunta_atual += " " + linha

    # Guardar última pergunta
    if pergunta_atual and opcoes_atuais:
        perguntas.append({
            "id": num_atual,
            "texto": pergunta_atual,
            "opcoes": opcoes_atuais
        })

    return perguntas



# =============================================================================
# FOLHA AZUL UCM — OMR exclusivo (gabarito FOLHA AZUL1.pdf)
# =============================================================================
from folha_azul.leitor import (
    FASES_OMR,
    PROXIMA_FASE_OMR,
    normalizar_fase_omr,
    processar_folha_azul,
)


def processar_folha_avaliacao(imagem_bgr, **kwargs):
    """Compatibilidade com testes antigos."""
    return processar_folha_azul(imagem_bgr, **kwargs)


def _sha256_bytes(payload: bytes) -> str:
    import hashlib
    return hashlib.sha256(payload).hexdigest()


def limpar_resultados_antigos(max_ficheiros: int = 40):
    pasta = Path("resultados")
    pasta.mkdir(exist_ok=True)
    ficheiros = sorted(
        [f for f in pasta.glob("*.json") if f.is_file()],
        key=lambda f: f.stat().st_mtime,
        reverse=True,
    )
    for f in ficheiros[max_ficheiros:]:
        try:
            f.unlink()
        except Exception:
            logger.warning("Falha ao remover resultado antigo: %s", f)


def guardar_json(resultado, prefixo="folha"):
    """
    Guarda um dicionário em resultados/{prefixo}_YYYYMMDD_HHMMSS.json.

    1. Cria a pasta resultados se não existir.
    2. Nome com data e hora.
    3. JSON formatado (utf-8).
    4. Devolve o caminho do ficheiro.
    """
    os.makedirs("resultados", exist_ok=True)
    ts = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    caminho = os.path.join("resultados", f"{prefixo}_{ts}.json")
    with open(caminho, "w", encoding="utf-8") as f:
        json.dump(resultado, f, ensure_ascii=False, indent=4)
    return caminho


# =============================================================================
# GUARDAR EM .DOCX
# =============================================================================

def guardar_docx(texto: str, nome: str) -> str:
    """Guarda o texto extraído num ficheiro Word e retorna o caminho."""
    os.makedirs("resultados", exist_ok=True)
    caminho = os.path.join("resultados", nome)

    doc = Document()
    titulo = doc.add_heading("Texto Extraído por OCR", 0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(
        f"Data: {datetime.datetime.now().strftime('%d/%m/%Y %H:%M:%S')}"
    ).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    doc.add_heading("Conteúdo:", 1)
    para = doc.add_paragraph(texto)
    para.runs[0].font.size = Pt(12)
    doc.save(caminho)
    return caminho


# =============================================================================
# ENDPOINTS
# =============================================================================

@app.get("/health")
async def health():
    """Verificar se o servidor está ativo."""
    return {
        "status": "ok",
        "servico": "OCR API",
        "versao": "1.0.0",
        "timestamp": datetime.datetime.now().isoformat()
    }


@app.post("/ocr/image")
async def ocr_image(file: UploadFile = File(...)):
    """
    Recebe uma imagem, extrai o texto com OCR e retorna JSON.

    Corpo da resposta:
    {
        "sucesso": true,
        "texto": "texto extraído completo",
        "palavras": 42,
        "linhas": ["linha 1", "linha 2", ...],
        "confianca": 0.94,
        "ficheiro_docx": "resultados/ocr_20260517_120000.docx",
        "metadados": {
            "nome_ficheiro": "foto.jpg",
            "tamanho_bytes": 12345,
            "timestamp": "2026-05-17T12:00:00"
        }
    }
    """
    # Verificar tipo de ficheiro
    if not file.content_type.startswith("image/"):
        raise HTTPException(400, "Ficheiro deve ser uma imagem")

    conteudo = await file.read()
    imagem = bytes_to_cv2(conteudo)

    if imagem is None:
        raise HTTPException(400, "Não foi possível processar a imagem")

    # Extrair texto
    texto = extrair_melhor_texto(imagem)

    if not texto:
        return JSONResponse({
            "sucesso": False,
            "texto": "",
            "palavras": 0,
            "linhas": [],
            "confianca": 0.0,
            "mensagem": "Nenhum texto detectado. Tenta com uma imagem mais nítida."
        })

    # Organizar linhas
    linhas = [l.strip() for l in texto.split('\n') if l.strip()]

    # Calcular confiança média
    try:
        processada = pre_processar(imagem, "contraste")
        img_pil    = Image.fromarray(processada)
        dados      = pytesseract.image_to_data(
            img_pil, output_type=pytesseract.Output.DICT, lang="por+eng"
        )
        confs = [int(c) for c in dados["conf"] if int(c) > 0]
        confianca = round(sum(confs) / len(confs) / 100, 2) if confs else 0.0
    except Exception:
        confianca = 0.0

    # Guardar .docx
    ts    = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    docx_path = guardar_docx(texto, f"ocr_{ts}.docx")

    return JSONResponse({
        "sucesso": True,
        "texto": texto,
        "palavras": len(texto.split()),
        "linhas": linhas,
        "confianca": confianca,
        "ficheiro_docx": docx_path,
        "metadados": {
            "nome_ficheiro": file.filename,
            "tamanho_bytes": len(conteudo),
            "timestamp": datetime.datetime.now().isoformat()
        }
    })


@app.post("/ocr/quiz")
async def ocr_quiz(file: UploadFile = File(...)):
    """
    Recebe imagem com perguntas de múltipla escolha (A/B/C/D),
    extrai e organiza em JSON estruturado.

    Corpo da resposta:
    {
        "sucesso": true,
        "total_perguntas": 3,
        "perguntas": [
            {
                "id": 1,
                "texto": "Qual é a capital de Moçambique?",
                "opcoes": {
                    "A": "Maputo",
                    "B": "Beira",
                    "C": "Nampula",
                    "D": "Tete"
                }
            }
        ],
        "texto_bruto": "...",
        "metadados": { ... }
    }
    """
    if not file.content_type.startswith("image/"):
        raise HTTPException(400, "Ficheiro deve ser uma imagem")

    conteudo = await file.read()
    imagem   = bytes_to_cv2(conteudo)

    if imagem is None:
        raise HTTPException(400, "Não foi possível processar a imagem")

    # Extrair texto
    texto = extrair_melhor_texto(imagem)

    if not texto:
        return JSONResponse({
            "sucesso": False,
            "total_perguntas": 0,
            "perguntas": [],
            "texto_bruto": "",
            "mensagem": "Nenhum texto detectado."
        })

    # Parsear quiz
    perguntas = parsear_quiz(texto)

    # Guardar .docx com perguntas formatadas
    ts = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    texto_formatado = ""
    for p in perguntas:
        texto_formatado += f"\n{p['id']}. {p['texto']}\n"
        for letra, opcao in p['opcoes'].items():
            texto_formatado += f"   {letra}) {opcao}\n"
    guardar_docx(texto_formatado or texto, f"quiz_{ts}.docx")

    return JSONResponse({
        "sucesso": True,
        "total_perguntas": len(perguntas),
        "perguntas": perguntas,
        "texto_bruto": texto,
        "metadados": {
            "nome_ficheiro": file.filename,
            "tamanho_bytes": len(conteudo),
            "timestamp": datetime.datetime.now().isoformat()
        }
    })


@app.post("/ocr/folha")
async def ocr_folha(
    file: UploadFile = File(...),
    fase: str = Form("codigo"),
    n_questoes_disciplina_1: int = Form(40),
    n_questoes_disciplina_2: int = Form(40),
):
    """
    Folha azul UCM (gabarito FOLHA AZUL1.pdf). Fotos P&B ou coloridas.

    - **fase=codigo** — grelha 10×10 do candidato (padrão)
    - **fase=disciplinas** — código + exame integrado (checkboxes Disc. 1 e 2)
    - **fase=opcoes** — + grelha respostas 1–80 (em calibração)
    """
    # 1. Validar se o ficheiro é imagem
    if file.content_type and not file.content_type.startswith("image/"):
        raise HTTPException(status_code=400, detail="Ficheiro deve ser uma imagem")

    conteudo = await file.read()
    request_id = str(uuid.uuid4())
    image_sha256 = _sha256_bytes(conteudo)
    logger.info(
        "OCR folha request_id=%s file='%s' bytes=%s sha256=%s",
        request_id,
        file.filename,
        len(conteudo),
        image_sha256[:16],
    )

    # 2. Ler a imagem com bytes_to_cv2()
    imagem = bytes_to_cv2(conteudo)
    if imagem is None:
        raise HTTPException(status_code=400, detail="Não foi possível processar a imagem")

    n1 = max(0, min(200, int(n_questoes_disciplina_1)))
    n2 = max(0, min(200, int(n_questoes_disciplina_2)))
    fase_norm = normalizar_fase_omr(fase)

    dados = processar_folha_azul(
        imagem,
        fase=fase_norm,
        n_disciplina_1=n1,
        n_disciplina_2=n2,
        image_sha256=image_sha256,
        request_id=request_id,
    )

    resultado: Dict[str, Any] = {
        "sucesso": True,
        "request_id": request_id,
        "imagem_sha256": image_sha256,
        "modelo_folha": dados.get("modelo_folha", "folha_azul_ucm"),
        "gabarito_versao": dados.get("gabarito_versao"),
        "fase": dados.get("fase", fase_norm),
        "proxima_fase": dados.get("proxima_fase"),
        "fases_disponiveis": dados.get("fases_disponiveis", list(FASES_OMR)),
        "codigo": dados.get("codigo", ""),
        "preparacao_imagem": dados.get("preparacao_imagem", {}),
        "parametros_pedido": {
            "fase": fase_norm,
            "n_questoes_disciplina_1": n1,
            "n_questoes_disciplina_2": n2,
        },
    }
    if dados.get("codigo_candidato") is not None:
        resultado["codigo_candidato"] = dados["codigo_candidato"]
    if dados.get("exame_integrado") is not None:
        resultado["exame_integrado"] = dados["exame_integrado"]
    if dados.get("respostas") is not None:
        resultado["respostas"] = dados["respostas"]
        resultado["respostas_marcadas"] = dados.get("respostas_marcadas", [])
        resultado["total_respostas"] = len(dados.get("respostas_marcadas", []))
        resultado["total_questoes"] = dados.get("total_questoes", 0)
    if dados.get("avisos"):
        resultado["avisos"] = dados["avisos"]

    # 7. Guardar em resultados/ e limpar histórico antigo
    resultado["ficheiro_json"] = guardar_json(resultado, prefixo="folha")
    limpar_resultados_antigos(max_ficheiros=40)

    # 8. Retornar para frontend / Postman
    return JSONResponse(
        resultado,
        headers={
            "Cache-Control": "no-store, no-cache, must-revalidate, max-age=0",
            "Pragma": "no-cache",
            "X-Request-Id": request_id,
        },
    )
